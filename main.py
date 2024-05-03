from tkinter import *
from tkinter import messagebox
from spacetype import Room, Apartment, MultistoryBuilding
import docx

def calculate_and_save_report():
    building_type = building_type_var.get()
    length = float(length_entry.get())
    width = float(width_entry.get())
    height = float(height_entry.get())
    
    if building_type == "Комната":
        building = Room(length, width, height)
        total_vol = building.calc_vol()
        heat_power = building.calc_heat_power()
    elif building_type == "Квартира":
        num_rooms = int(rooms_entry.get())
        building = Apartment(length, width, height, num_rooms)
        total_vol = building.calc_total_vol()
        heat_power = building.calc_heat_power()
    else:  # Многоэтажный дом
        num_floors = int(floors_entry.get())
        num_units_per_floor = int(units_entry.get())
        building = MultistoryBuilding(length, width, height, num_floors, num_units_per_floor)
        total_vol = building.calc_total_vol()
        heat_power = building.calc_heat_power()

    # Отображение результатов
    result_label.config(text=f"Общий объем: {total_vol} кв.м\nТепловая мощность: {heat_power} Вт")

    # Сохранение результатов в отчет .docx
    doc = docx.Document()
    doc.add_heading('Результаты расчетов', level=1)
    doc.add_paragraph(f"Общая площадь: {total_vol} кв.м")
    doc.add_paragraph(f"Тепловая мощность: {heat_power} Вт")
    doc.save('report.docx')
    
    messagebox.showinfo("Сохранение", "Результаты сохранены в отчет.docx ")

# Создание GUI
root = Tk()
root.title("Калькулятор строительства")

building_type_var = StringVar()
building_type_var.set("Комната")

building_type_label = Label(root, text="Выберите тип помещения:")
building_type_label.pack()

building_type_options = ["Комната", "Квартира", "Многоэтажный дом"]
building_type_menu = OptionMenu(root, building_type_var, *building_type_options)
building_type_menu.pack()

length_label = Label(root, text="Длина:")
length_label.pack()
length_entry = Entry(root)
length_entry.pack()

width_label = Label(root, text="Ширина:")
width_label.pack()
width_entry = Entry(root)
width_entry.pack()

height_label = Label(root, text="Высота:")
height_label.pack()
height_entry = Entry(root)
height_entry.pack()

rooms_label = Label(root, text="Количество комнат (для квартиры):")
rooms_label.pack()
rooms_entry = Entry(root)
rooms_entry.pack()

floors_label = Label(root, text="Этажей (для многоэтажного дома):")
floors_label.pack()
floors_entry = Entry(root)
floors_entry.pack()

units_label = Label(root, text="Количество квартир на этаже (для многоэтажного дома):")
units_label.pack()
units_entry = Entry(root)
units_entry.pack()

calculate_button = Button(root, text="Рассчитать и сохранить", command=calculate_and_save_report)
calculate_button.pack()

result_label = Label(root, text="")
result_label.pack()

root.mainloop()
